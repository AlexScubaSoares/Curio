from docx import Document   
from docx.opc.exceptions import PackageNotFoundError
import pyodbc
import os
import time

SERVER = 'XANDAOTRAVEL\SQLEXPRESS'
DATABASE = 'CURIO_DB'

connectionString = f'DRIVER={{SQL Server}};SERVER={SERVER};DATABASE={DATABASE};Trusted_Connection=True'
conn = pyodbc.connect(connectionString)
cursor = conn.cursor()

def trata_arquivo_db(arquivo: str, criacao: str):
    sql_arq = f"SELECT count(*) as QTDE FROM ARQUIVOS WHERE NOME_ARQUIVO = ?;"

    cursor.execute(
        sql_arq,
        arquivo
        )
    
    retorno = cursor.fetchall()

    for ret in retorno:
        print('Dentro do for ret ',arquivo, '==> ',ret[0])

    if ret[0] == 0:
        sql_insert_arq = f"INSERT INTO ARQUIVOS(NOME_ARQUIVO, CRIACAO) VALUES(?,?)"
        cursor.execute(
            sql_insert_arq,
            arquivo,
            criacao
        )

        conn.commit()    

    sql_arq = f"SELECT COD_ARQUIVO FROM ARQUIVOS WHERE NOME_ARQUIVO = ?;"

    cursor.execute(
        sql_arq,
        arquivo
        )
    
    retorno = cursor.fetchall()
    for cod in retorno:
        print(cod[0])
    print('Vai retornar o codigo do arquivo ', arquivo, ' ==> ', cod[0])
    return cod[0]

def grava_atividade_db(cod_arq: int, texto:str):
    sql_insert = 'INSERT INTO conteudos(cod_arquivo, atividade) values(?,?);'
    cursor.execute(sql_insert,
                   cod_arq,
                   texto)
    conn.commit()

class ReadDocxFile:
    def __init__(self, filename):
        self.fname = filename
        self.loaded_document = None
        self.content = None

    def load_document(self):
        try:
            self.loaded_document = Document(self.fname)
            self.content = self.get_content_as_list()
        except PackageNotFoundError as e:
            print(f'Arquivo não encontrado: {self.fname}')
            quit()

    def get_cleaned_text(self):
        texto_limpo = []
        nada = 0
        for paragraph in self.loaded_document.paragraphs:
            if paragraph.text == '':
                nada += 1
            else:
                texto_limpo.append(paragraph.text.replace("\t", ""))
        return texto_limpo
    
    def get_content_as_list(self):
        return [paragraph.text for paragraph in self.loaded_document.paragraphs]

def files_path(path):
    for p, _, files in os.walk(os.path.abspath(path)):
        for file in files:
            arq_work = os.path.join(p, file)
            print(f'abrindo o arquivo: {arq_work}')

            if os.path.isfile(arq_work) and file.lower().endswith(".docx") and file.lower().startswith("material_"):
               docx = ReadDocxFile(arq_work)
               docx.load_document()
               texto_final = docx.get_cleaned_text()
               dtcria = os.path.getmtime(arq_work)
               criadt = time.ctime(dtcria)
               dtformated = time.strptime(criadt)
               criacao = time.strftime("%Y-%m-%d", dtformated)
               cod_ret = trata_arquivo_db(arq_work, criacao)
               for i in range(len(texto_final)):
                    grava_atividade_db(cod_ret, texto_final[i])
                    print(p, ' --- ', file, cod_ret)

if __name__ == "__main__":

    files_path("/OLGA")

    sql_query = "SELECT arq.nome_arquivo, cont.atividade FROM ARQUIVOS arq, CONTEUDOS cont WHERE cont.COD_ARQUIVO = arq.COD_ARQUIVO;"
    cursor.execute(sql_query)
    dados = cursor.fetchall()
    for d in dados:
        print(f'{d.nome_arquivo} --- {d.atividade}')
    

cursor.close
conn.close
